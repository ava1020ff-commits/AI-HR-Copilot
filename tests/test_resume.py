"""只使用内存合成文档、临时数据库和模拟 API。"""

import io
import json
import sqlite3
from pathlib import Path
from unittest.mock import MagicMock

import pytest
import requests
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
from streamlit.testing.v1 import AppTest

from database.candidates import save_candidate
from services.jd_parser import LLMConfig
from services.resume_parser import FIELDS, ResumeError, extract_resume, matching_fields, parse_resume, validate_resume

TEXT = """姓名: 测试甲
性别: 女
年龄: 28 岁
婚姻状况: 未婚
教育经历
测试大学 / 本科 / 计算机 / 2018-2022
工作经历
测试公司 / Python 工程师 / 2022-2024
实习经历
实习公司 / 开发实习生 / 2021
项目经历
招聘系统 / 后端开发 / 实现检索 API
技能: Python, SQL
证书: 测试认证
"""
PAGE = Path(__file__).resolve().parents[1] / "pages" / "02_候选人.py"


@pytest.fixture(autouse=True)
def isolation(monkeypatch, tmp_path):
    monkeypatch.setenv("RESUME_DATABASE_PATH", str(tmp_path / "candidates.sqlite3"))
    for name in ("LLM_API_KEY", "LLM_MODEL", "LLM_BASE_URL"):
        monkeypatch.delenv(name, raising=False)
    monkeypatch.setattr("services.resume_parser.requests.post", MagicMock(side_effect=AssertionError("禁止真实 API")))


def docx_bytes(text: str = TEXT, table: bool = False) -> bytes:
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    if table:
        cell = document.add_table(rows=1, cols=1).cell(0, 0)
        cell.text = "表格技能 Python"
        cell.add_table(rows=1, cols=1).cell(0, 0).text = "嵌套 SQL"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def pdf_bytes(blank: bool = False, encrypted: bool = False) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if not blank:
        font = DictionaryObject({NameObject("/Type"): NameObject("/Font"), NameObject("/Subtype"): NameObject("/Type1"), NameObject("/BaseFont"): NameObject("/Helvetica")})
        page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})})
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 50 740 Td (Name: Test Candidate) Tj 0 -20 Td (Skills: Python, SQL) Tj ET")
        page[NameObject("/Contents")] = writer._add_object(stream)
    if encrypted:
        writer.encrypt("test-only-password")
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def test_extract_docx_and_nested_tables() -> None:
    text = extract_resume("resume.DOCX", docx_bytes(table=True))
    assert "测试甲" in text and "表格技能 Python" in text and "嵌套 SQL" in text


def test_pdf_text_and_parse() -> None:
    text = extract_resume("resume.pdf", pdf_bytes())
    result = parse_resume(text)
    assert result["candidate_name"] == "Test Candidate"
    assert result["skills"] == ["Python", "SQL"]


@pytest.mark.parametrize("filename,content", [("a.txt", b"text"), ("a.pdf", b""), ("a.pdf", b"fake"), ("a.docx", b"fake"), ("a.pdf", b"x" * (10 * 1024 * 1024 + 1))], ids=["extension", "empty", "bad_pdf", "bad_docx", "oversize"])
def test_bad_files(filename: str, content: bytes) -> None:
    with pytest.raises(ResumeError):
        extract_resume(filename, content)


def test_scanned_blank_and_encrypted_pdf() -> None:
    with pytest.raises(ResumeError, match="OCR"):
        extract_resume("blank.pdf", pdf_bytes(blank=True))
    with pytest.raises(ResumeError, match="加密"):
        extract_resume("secret.pdf", pdf_bytes(encrypted=True))


def test_empty_docx() -> None:
    with pytest.raises(ResumeError, match="文字"):
        extract_resume("empty.docx", docx_bytes(""))


def test_local_fields_and_privacy() -> None:
    result = parse_resume(TEXT)
    assert result["candidate_name"] == "测试甲"
    assert all(result[key] for key in FIELDS)
    assert "实习公司" in result["internships"][0]
    assert result["skills"] == ["Python", "SQL"]
    encoded = json.dumps(result, ensure_ascii=False)
    assert not any(value in encoded for value in ("性别", "28", "未婚"))
    assert set(matching_fields(result)) == set(FIELDS)
    assert "candidate_name" not in matching_fields(result)


@pytest.mark.parametrize("key,value", [("gender", "女"), ("age", 28), ("photo", "url"), ("marital_status", "未婚")])
def test_forbidden_extra_fields(key: str, value: object) -> None:
    result = parse_resume(TEXT)
    result[key] = value
    with pytest.raises(ResumeError):
        validate_resume(result)


@pytest.mark.parametrize("value", ["性别: 女", "年龄: 28", "婚姻: 已婚", "照片: url", "female", "28 years old"])
def test_sensitive_values_rejected(value: str) -> None:
    result = parse_resume(TEXT)
    result["skills"].append(value)
    with pytest.raises(ResumeError):
        validate_resume(result)


def test_normalization_and_required_name() -> None:
    result = parse_resume(TEXT)
    result["skills"] = [" Ｐｙｔｈｏｎ ", "Python", ""]
    assert validate_resume(result)["skills"] == ["Python"]
    result["candidate_name"] = ""
    with pytest.raises(ResumeError):
        validate_resume(result, require_name=True)


def test_no_save_without_confirmation(tmp_path) -> None:
    path = tmp_path / "not_saved.sqlite3"
    with pytest.raises(ValueError):
        save_candidate(parse_resume(TEXT), "local", db_path=path)
    assert not path.exists()


def test_sqlite_edited_data_and_dedup(tmp_path) -> None:
    result = parse_resume(TEXT)
    result["skills"] = ["人工修改后的 SQL"]
    path = tmp_path / "saved.sqlite3"
    first = save_candidate(result, "local", confirmed=True, db_path=path)
    assert first == save_candidate(result, "local", confirmed=True, db_path=path)
    with sqlite3.connect(path) as connection:
        row = connection.execute("SELECT parsed_json FROM candidates").fetchone()
        assert json.loads(row[0]) == result
        assert connection.execute("SELECT COUNT(*) FROM candidates").fetchone()[0] == 1
        columns = {row[1] for row in connection.execute("PRAGMA table_info(candidates)")}
        assert not columns.intersection({"source_text", "photo", "gender", "age", "marital_status"})


def test_llm_request_and_invalid_response(monkeypatch) -> None:
    response = MagicMock()
    response.__enter__.return_value = response
    response.status_code = 200
    response.json.return_value = {"choices": [{"finish_reason": "stop", "message": {"content": json.dumps(parse_resume(TEXT))}}]}
    post = MagicMock(return_value=response)
    monkeypatch.setattr("services.resume_parser.requests.post", post)
    config = LLMConfig("test-key", "https://example.invalid/v1", "test-model")
    assert parse_resume(TEXT, False, config)["candidate_name"] == "测试甲"
    payload = post.call_args.kwargs["json"]
    assert "性别" not in payload["messages"][1]["content"]
    assert payload["response_format"] == {"type": "json_object"}
    response.json.return_value = {"choices": []}
    with pytest.raises(ResumeError):
        parse_resume(TEXT, False, config)
    post.side_effect = requests.Timeout("private")
    with pytest.raises(ResumeError, match="超时"):
        parse_resume(TEXT, False, config)


def uploaded_docx() -> io.BytesIO:
    upload = io.BytesIO(docx_bytes())
    upload.name = "synthetic.docx"
    return upload


def test_page_edit_and_confirm(monkeypatch, tmp_path) -> None:
    upload = uploaded_docx()
    monkeypatch.setattr("streamlit.file_uploader", lambda *args, **kwargs: upload)
    app = AppTest.from_file(str(PAGE.parents[1] / "app.py"), default_timeout=15).run().switch_page("pages/" + PAGE.name).run()
    app.button[0].click().run()
    assert not app.exception
    path = tmp_path / "candidates.sqlite3"
    assert not path.exists()
    app.text_input[0].input("测试乙")
    app.text_area(key="resume_edit_skills").input("SQL\nGit")
    app.button[1].click().run()
    assert app.error and not path.exists()
    app.checkbox(key="resume_reviewed").check()
    app.button[1].click().run()
    assert not app.exception and app.success
    with sqlite3.connect(path) as connection:
        saved = json.loads(connection.execute("SELECT parsed_json FROM candidates").fetchone()[0])
        assert saved["candidate_name"] == "测试乙"
        assert saved["skills"] == ["SQL", "Git"]


def test_removing_upload_clears_draft(monkeypatch) -> None:
    holder = {"upload": uploaded_docx()}
    monkeypatch.setattr("streamlit.file_uploader", lambda *args, **kwargs: holder["upload"])
    app = AppTest.from_file(str(PAGE.parents[1] / "app.py"), default_timeout=15).run().switch_page("pages/" + PAGE.name).run()
    app.button[0].click().run()
    assert app.text_input
    holder["upload"] = None
    app.run()
    assert not app.exception and not app.text_input
    assert "resume_draft" not in app.session_state


def test_page_navigation() -> None:
    app = AppTest.from_file(str(PAGE.parents[1] / "app.py"), default_timeout=15).run()
    app.switch_page("pages/02_候选人.py").run()
    assert not app.exception
    assert app.title[0].value == "候选人"


def test_save_failure_preserves_edits(monkeypatch) -> None:
    upload = uploaded_docx()
    monkeypatch.setattr("streamlit.file_uploader", lambda *args, **kwargs: upload)
    monkeypatch.setattr("database.candidates.save_candidate", MagicMock(side_effect=sqlite3.OperationalError("locked")))
    app = AppTest.from_file(str(PAGE.parents[1] / "app.py"), default_timeout=15).run().switch_page("pages/" + PAGE.name).run()
    app.button[0].click().run()
    app.text_input[0].input("修改保留")
    app.checkbox(key="resume_reviewed").check()
    app.button[1].click().run()
    assert not app.exception and app.error and not app.success
    assert app.text_input[0].value == "修改保留"


def test_same_filename_new_content_clears_draft(monkeypatch) -> None:
    holder = {"upload": uploaded_docx()}
    monkeypatch.setattr("streamlit.file_uploader", lambda *args, **kwargs: holder["upload"])
    app = AppTest.from_file(str(PAGE.parents[1] / "app.py"), default_timeout=15).run().switch_page("pages/" + PAGE.name).run()
    app.button[0].click().run()
    new_upload = io.BytesIO(docx_bytes(TEXT.replace("测试甲", "测试乙")))
    new_upload.name = "synthetic.docx"
    holder["upload"] = new_upload
    app.run()
    assert not app.text_input
    app.button[0].click().run()
    assert app.text_input[0].value == "测试乙"


def test_llm_consent_required(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("LLM_API_KEY", "test-only-key")
    monkeypatch.setenv("LLM_MODEL", "test-model")
    upload = uploaded_docx()
    monkeypatch.setattr("streamlit.file_uploader", lambda *args, **kwargs: upload)
    app = AppTest.from_file(str(PAGE.parents[1] / "app.py"), default_timeout=15).run().switch_page("pages/" + PAGE.name).run()
    app.button[0].click().run()
    assert not app.exception and app.error
    assert "授权" in app.error[0].value
    assert not (tmp_path / "candidates.sqlite3").exists()


def test_invalid_edited_sensitive_data_not_saved(tmp_path) -> None:
    data = parse_resume(TEXT)
    data["work_experience"].append("婚姻状况: 已婚")
    path = tmp_path / "sensitive.sqlite3"
    with pytest.raises(ResumeError):
        save_candidate(data, "local", confirmed=True, db_path=path)
    assert not path.exists()
